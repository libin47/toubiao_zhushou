import io
import os
import re
import unicodedata
from PIL import Image
from docx import Document
from docx.enum.table import WD_TABLE_DIRECTION, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.text.paragraph import Paragraph
from docx.styles.style import ParagraphStyle, CharacterStyle

from app.hidden_clean_config import HiddenCleanerConfig, ConfigParagraph, ConfigFont, ConfigImage, ConfigBase
# config = yaml.safe_load(open('config.yml', 'r', encoding='utf-8'))

fmt_dict = {
    "一 二 三": "chineseCounting",
    "1 2 3": "decimal",
    "１　２　３": "decimalFullWidth",
    "壹 贰 叁": "chineseLegalSimplified",
    "① ② ③": "decimalEnclosedCircleChinese",
    "Ⅰ Ⅱ Ⅲ": "upperRoman",
    "甲 乙 丙": "ideographTraditional",
    "A B C": "upperLetter"
}
fmt_index = [0,0,0,0,0,0,0,0,0]

def add_xml(p, name, value=""):
    w = OxmlElement(name)
    # value为空时，不添加子属性
    if value=="":
        pass
    # value为字符串或字典时时，添加子属性
    elif type(value)==dict:
        for key, val in value.items():
            w.set(qn(key), str(val))
    else:
        w.set(qn('w:val'), str(value))
    p.append(w)
    return w

def convert_image_to_grayscale(image_bytes, config: ConfigBase):
    # 将字节数据转换为PIL图像
    image = Image.open(io.BytesIO(image_bytes))
    # 缩放
    width, height = image.size
    if config.图片缩放.get() and width>config.图片最大宽度.get():
        new_width = config.图片最大宽度.get()
        new_height = int(new_width * (height / width))
        image = image.resize((new_width, new_height))
    # 转换为灰度图
    grayscale_image = image.convert("L")
    # 保存为字节数据
    output = io.BytesIO()
    grayscale_image.save(output, format="PNG")
    return output.getvalue()


def insert_image(image_bytes, run):
    # 将字节数据转换为PIL图像
    image_stream = io.BytesIO(image_bytes)
    image = Image.open(image_stream)
    # 缩放
    width, height = image.size
    hcm = 15 * height / width
    # 插入
    # run.add_picture(image_stream, width=width, height=height)
    run.add_picture(image_stream, width=Cm(15), height=Cm(hcm))
    return True

def check_docx(file):
    # 检查docx文件是否存在
    if not os.path.exists(file):
        return False
    # 检查docx文件是否为空
    if os.path.getsize(file) == 0:
        return False
    return True


def is_image(paragraph):
    for run in paragraph.runs:
        if run._element.xpath('.//a:blip'):
            return True
    return False

def check_char(char, config:ConfigBase):
    """
    检查字符是否是特殊字符
    """
    if char in config.特殊字符_保留.get():
        return True
    if char in config.特殊字符_删除.get():
        return False
    # 基本希腊字母
    if 0x0020<= ord(char) <= 0x007e:
        return True
    # 中日韩部首补充
    if 0x2E80<= ord(char) <= 0x2EFF:
        return True
    # 康熙部首
    if 0x2F00<= ord(char) <= 0x2FDF:
        return True
    # 中日韩符号和标点
    if 0x3000<= ord(char) <= 0x303F:
        return True
    # 中日韩统一表意文字扩展区A
    if 0x3400<= ord(char) <= 0x4DBF:
        return True
    # 中日韩统一表意文字
    if 0x4E00<= ord(char) <= 0x9FFF:
        return True
    # 中日韩兼容表意文字
    if 0xF900<= ord(char) <= 0xFAFF:
        return True
    # 半角及全角形式
    if 0xFF00 <= ord(char) <= 0xFFEF:
        return True
    # 中日韩统一表意文字扩展区B
    if 0x20000 <= ord(char) <= 0x2A6DF:
        return True
    # 中日韩统一表意文字扩展区C D E F I
    if 0x2A700 <= ord(char) <= 0x2EE5F:
        return True
    # 中日韩统一表意文字扩展区G H
    if 0x30000 <= ord(char) <= 0x323AF:
        return True
    return False

def set_super_char(sentence, config:ConfigBase):
    # 特殊字符
    if config.删除特殊字符.get():
        result = ""
        for char in sentence:
            if check_char(char, config):
                result += char
        sentence = result
    return sentence

def set_char(sentence, config:ConfigBase):
    # 英文标点to中文标点
    if config.标点转中文.get():
        for i in range(len(config.中英文标点字典[0])):
            sentence = sentence.replace(config.中英文标点字典[0][i], config.中英文标点字典[1][i])
    # 半角2全角
    if config.半角转为全角.get():
        full_width_text = ""
        for char in sentence:
            if 0x0020 <= ord(char) <= 0x007e:  # 半角字符范围
                full_width_char = unicodedata.lookup('FULLWIDTH ' + unicodedata.name(char))
                full_width_text += full_width_char
            else:
                full_width_text += char
        sentence = full_width_text
    return sentence


def set_row(row, config: HiddenCleanerConfig):
    # 删除边框样式
    # for ele in row._element.xpath('.//w:tblBorders'):
    #     ele.getparent().remove(ele)
    # 设置行高
    if config.table.style.行高方式.get()=="自适应":
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
    elif config.table.style.行高方式.get()=="固定":
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(config.table.style.行高.get())
    elif config.table.style.行高方式.get()=="最小值":
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(config.table.style.行高.get())
    else:
        print("行高方式错误，请检查配置文件: %s"%config.table.style.行高方式.get())


def set_cell(cell, config:HiddenCleanerConfig):
    # 删除边框样式
    # for ele in cell._element.xpath('.//w:tcBorders'):
    #     ele.getparent().remove(ele)
    # 垂直对齐
    if config.table.style.垂直对齐.get() == "居中":
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    elif config.table.style.垂直对齐.get() == "底部对齐":
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    elif config.table.style.垂直对齐.get() == "顶部对齐":
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    else:
        print("垂直对齐方式错误，请检查配置文件: %s"%config.table.style.垂直对齐.get())


def set_table(table, config:HiddenCleanerConfig):
    # 表格对齐
    if config.table.style.对齐.get() == "左对齐":
        table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif config.table.style.对齐.get() == "右对齐":
        table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif config.table.style.对齐.get() == "居中":
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        print("表格对齐方式错误，请检查配置文件: %s"%config.table.style.对齐.get())
    # 方向
    if config.table.style.表格方向.get() == "从左到右":
        table.table_direction = WD_TABLE_DIRECTION.LTR
    elif config.table.style.表格方向.get() == "从右到左":
        table.table_direction = WD_TABLE_DIRECTION.RTL
    else:
        print("表格方向错误，请检查配置文件: %s"%config.table.style.表格方向.get())
    # 自动调整列宽
    table.autofit = config.table.style.自动调整列宽.get()
    # 边框
    for ele in table._element.xpath('.//w:tcBorders'):
        ele.getparent().remove(ele)
    for ele in table._element.xpath('.//w:tblBorders'):
        ele.getparent().remove(ele)
    border = add_xml(table._element.tblPr, 'w:tblBorders', "")
    if config.table.style.边框颜色.get()=="黑色":
        color = RGBColor(0,0,0)
    elif config.table.style.边框颜色.get()=="红色":
        color = RGBColor(255,0,0)
    else:
        color = RGBColor(0,0,0)
    add_xml(border, 'w:top', {"w:val": config.table.style.边框.get(), "w:sz": config.table.style.边框粗细.get(), "w:space": "0", "w:color": color})
    add_xml(border, 'w:left', {"w:val": config.table.style.边框.get(), "w:sz": config.table.style.边框粗细.get(), "w:space": "0", "w:color": color})
    add_xml(border, 'w:bottom', {"w:val": config.table.style.边框.get(), "w:sz": config.table.style.边框粗细.get(), "w:space": "0", "w:color": color})
    add_xml(border, 'w:right', {"w:val": config.table.style.边框.get(), "w:sz": config.table.style.边框粗细.get(), "w:space": "0", "w:color": color})
    add_xml(border, 'w:insideH', {"w:val": config.table.style.边框.get(), "w:sz": config.table.style.边框粗细.get(), "w:space": "0", "w:color": color})
    add_xml(border, 'w:insideV', {"w:val": config.table.style.边框.get(), "w:sz": config.table.style.边框粗细.get(), "w:space": "0", "w:color": color})

def remove_w14_styles(paragraph):
    for ele in paragraph._element.xpath('.//w14:*'):
        ele.getparent().remove(ele)

def set_paragraph(paragraph, config:HiddenCleanerConfig, table=False):
    if table:
        config_image = config.table.image
        config_paragraph = config.table.paragraph
    else:
        config_image = config.main.image
        config_paragraph = config.main.paragraph
    # 清除w14样式
    if config.extend.清除w14样式.get():
        remove_w14_styles(paragraph)
    # 删除自动编号
    if config.extend.删除自动编号.get():
        for pPr in paragraph._element.iter(qn('w:numPr')):
            pPr.getparent().remove(pPr)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    image_paragraph = is_image(paragraph)
    if image_paragraph:
        # 首行缩进2字符
        paragraph_format.first_line_indent = 0
        paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), str(int(config_image.首行缩进.get()) * 100))
        # 对齐方式
        if config_image.对齐方式.get() == "左对齐":
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 对齐方向
        elif config_image.对齐方式.get()  == "右对齐":
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 对齐方向
        elif config_image.对齐方式.get()  == "居中":
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 对齐方向
        # 行距
        行距方式 = config_image.行距方式.get()
        行距 = config_image.行距.get()
    else:
        # 首行缩进2字符
        paragraph_format.first_line_indent = 0
        paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), str(int(config_paragraph.首行缩进.get())*100))
        # 对齐方式
        if config_paragraph.对齐方式.get() =="左对齐":
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 对齐方向
        elif config_paragraph.对齐方式.get()=="右对齐":
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 对齐方向
        elif config_paragraph.对齐方式.get()=="居中":
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 对齐方向
        # 行距
        行距方式 = config_paragraph.行距方式.get()
        行距 = config_paragraph.行距.get()
    # 对齐到网络
    add_xml(paragraph_format.element.pPr, 'w:snapToGrid', str(config_paragraph.对齐网络.get()))
    # 右对齐网络
    add_xml(paragraph_format.element.pPr, 'w:adjustRightInd', str(config_paragraph.右对齐网络.get()))
    # 行距/行高
    if 行距方式 == "倍率":
        if 行距 == 1.5:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif 行距 == 2:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        elif 行距 == 1:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        elif 行距 == 0:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        else:
            paragraph_format.line_spacing = 行距
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    elif 行距方式 == "固定":
        paragraph_format.line_spacing = Pt(行距)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    # 孤行控制
    paragraph_format.widow_control = True if config_paragraph.孤行控制.get() else False
    # 默认值
    paragraph_format.left_indent = Pt(0) # 左缩进
    paragraph_format.right_indent = Pt(0) # 右缩进
    paragraph_format.space_before = Pt(0) # 段前间距
    paragraph_format.space_after = Pt(0)  # 段后间距
    paragraph_format.keep_together = False # 避免段落被拆分到下一页
    paragraph_format.keep_with_next = False # 段落和下一段保持同一页
    paragraph_format.page_break_before = False # 段落应显示在页面顶部
    paragraph_format.tab_stops.clear_all() # 清除所有制表位
    # 删除空行
    if config.base.删除空行.get() and not image_paragraph:
        if not paragraph.text.strip():
            p = paragraph._element
            p.getparent().remove(p)

def set_font(run, cfg_font:ConfigFont, cfg_base: ConfigBase):
    if run._element.xpath('.//a:blip'):
        set_image(run, cfg_base)
        return
    # 设置字体样式
    # 字体
    fontname = cfg_font.字体.get()
    run.font.name = fontname
    run._element.rPr.rFonts.set(qn('w:ascii'), fontname)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), fontname)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    # 字号
    run.font.size = Pt(cfg_font.字号.get())
    # 颜色
    if cfg_font.颜色.get()=="黑色":
        run.font.color.rgb = RGBColor(0,0,0)
    elif cfg_font.颜色.get()=="红色":
        run.font.color.rgb = RGBColor(255,0,0)
    # 对齐到网络
    run.font.snap_to_grid =  True if cfg_font.对齐到网络.get() else False
    # 字符缩放
    add_xml(run._element.rPr, "w:w", cfg_font.字符缩放.get())
    # 字符间距
    add_xml(run._element.rPr, "w:spacing", 0)
    # 字符间距调整
    add_xml(run._element.rPr, "w:kern", 0)
    # 其他
    run.font.italic = False # 斜体
    run.font.bold = False # 加粗
    run.font.underline = False # 下划线
    run.font.all_caps = False # 全大写
    run.font.double_strike = False # 双删除线
    run.font.strike = False # 删除线
    run.font.subscript = False # 下标
    run.font.superscript = False # 上标
    run.font.outline = False # 描边
    run.font.shadow = False # 阴影
    run.font.small_caps = False # 小大写
    run.font.emboss = False # 浮雕
    run.font.complex_script = False # 复杂语种
    run.font.cs_bold = False
    run.font.cs_italic = False
    if cfg_font.高亮.get():
        run.font.highlight_color = RGBColor(255, 255, 0)  # 设置高亮颜色为黄色
    else:
        run.font.highlight_color = None
    if cfg_base.删除空格.get():
        run.text = run.text.replace(' ', '')
    if cfg_base.删除制表符.get():
        run.text = run.text.replace('\t', '')  # 删除制表符
    if cfg_base.删除特殊字符.get():
        run.text = set_super_char(run.text, cfg_base)
    run.text = set_char(run.text, cfg_base)

def set_image(run, config:ConfigBase):
    if config.删除图片.get():
        run._element.getparent().remove(run._element)
    else:
        if config.图片灰度化.get():
            blip_elements = run._element.xpath('.//a:blip')
            if len(blip_elements)>1:
                # 目前未发现此种情况，先按只有一个处理
                print("图片数量大于1!!")
            for blip in blip_elements:
                # 获取图片数据
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                part = run.part.related_parts[embed]
                try:
                    new_image = convert_image_to_grayscale(part.image.blob, config)
                except:
                    print("图片处理错误，尝试使用源码转换")
                    new_image = convert_image_to_grayscale(part.blob, config)
                part._blob = new_image
        if config.图片嵌入.get():
            # 如果图片嵌入方式是anchor
            anchors = run._element.xpath('.//wp:anchor')
            if len(anchors) > 0:
                # 1、保存原来图片
                blip_elements = run._element.xpath('.//a:blip')
                images = []
                for blip in blip_elements:
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    part = run.part.related_parts[embed]
                    image_stream = part.blob
                    images.append(image_stream)
                # 2、删除原来图片
                for anchor in anchors:
                    anchor.getparent().remove(anchor)
                # 3、插入新图片
                for image_stream in images:
                    result = insert_image(image_stream, run)


def set_style(style, config:HiddenCleanerConfig):
    cfg_font = config.main.font
    cfg_paragraph = config.main.paragraph
    # 编号
    try:
        if config.extend.设置标题编号.get():
            if style.element.pPr is not None and style.element.pPr.find(qn("w:numPr")):
                try:
                    numPr = style.element.pPr.find(qn("w:numPr"))
                    # 设置numID
                    numId = numPr.find(qn("w:numId"))
                    numId.set(qn("w:val"), str(config.extend.标题编号ID.get()))
                    # 设置ilvl与大纲级别一致
                    outlineLv = style.element.pPr.find(qn("w:outlineLvl"))
                    lvl = outlineLv.get(qn("w:val"))
                    ilvl = numPr.find(qn("w:ilvl"))
                    ilvl.set(qn("w:val"), lvl)
                except:
                    pass
    except:
        pass
    # 字体
    try:
        fontname = cfg_font.字体.get()
        # TODO:可能会报错，原因未知
        style.font.name = fontname
        style._element.rPr.rFonts.set(qn('w:ascii'), fontname)
        style._element.rPr.rFonts.set(qn('w:hAnsi'), fontname)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
        style.font.size = Pt(cfg_font.字号.get())
    except:
        pass
    # 颜色
    try:
        if cfg_font.颜色.get() == "黑色":
            style.font.color.rgb = RGBColor(0, 0, 0)
        elif cfg_font.颜色.get() == "红色":
            style.font.color.rgb = RGBColor(255, 0, 0)
    except:
        pass
    # 对齐到网络
    try:
        style.font.snap_to_grid = True if cfg_font.对齐到网络.get() else False
    except:
        pass
    # 其他
    try:
        style.font.italic = False  # 斜体
        style.font.bold = False  # 加粗
        style.font.underline = False  # 下划线
        style.font.all_caps = False  # 全大写
        style.font.double_strike = False  # 双删除线
        style.font.strike = False  # 删除线
        style.font.subscript = False  # 下标
        style.font.superscript = False  # 上标
        style.font.outline = False  # 描边
        style.font.shadow = False  # 阴影
        style.font.small_caps = False  # 小大写
        style.font.emboss = False  # 浮雕
        style.font.complex_script = False  # 复杂语种
        style.font.cs_bold = False
        style.font.cs_italic = False
        if cfg_font.高亮.get():
            style.font.highlight_color = RGBColor(255, 255, 0)  # 设置高亮颜色为黄色
        else:
            style.font.highlight_color = None
        if type(style) == CharacterStyle:
            return
    except:
        pass
    # 段落
    # 首行缩进2字符
    try:
        style.paragraph_format.first_line_indent = 0
        style.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), str(int(cfg_paragraph.首行缩进.get())*100))
    except:
        pass
    # 对齐方式
    try:
        if cfg_paragraph.对齐方式.get() =="左对齐":
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 对齐方向
        elif cfg_paragraph.对齐方式.get()=="右对齐":
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 对齐方向
        elif cfg_paragraph.对齐方式.get()=="居中":
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 对齐方向
    except:
        pass
    # 行距
    try:
        行距方式 = cfg_paragraph.行距方式.get()
        行距 = cfg_paragraph.行距.get()
    except:
        pass
    try:
        # 对齐到网络
        add_xml(style.paragraph_format.element.pPr, 'w:snapToGrid', str(cfg_paragraph.对齐网络.get()))
        # 右对齐网络
        add_xml(style.paragraph_format.element.pPr, 'w:adjustRightInd', str(cfg_paragraph.右对齐网络.get()))
        # 行距/行高
        if 行距方式 == "倍率":
            if 行距 == 1.5:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif 行距 == 2:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            elif 行距 == 1:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            elif 行距 == 0:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            else:
                style.paragraph_format.line_spacing = 行距
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        elif 行距方式 == "固定":
            style.paragraph_format.line_spacing = Pt(行距)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    except:
        pass
    try:
        # 孤行控制
        style.paragraph_format.widow_control = True if cfg_paragraph.孤行控制.get() else False
        # 默认值
        style.paragraph_format.left_indent = Pt(0) # 左缩进
        style.paragraph_format.right_indent = Pt(0) # 右缩进
        style.paragraph_format.space_before = Pt(0) # 段前间距
        style.paragraph_format.space_after = Pt(0)  # 段后间距
        style.paragraph_format.keep_together = False # 避免段落被拆分到下一页
        style.paragraph_format.keep_with_next = False # 段落和下一段保持同一页
        style.paragraph_format.page_break_before = False # 段落应显示在页面顶部
        style.paragraph_format.tab_stops.clear_all() # 清除所有制表位
    except:
        pass



def get_abstract_num(abs_list):
    for i in range(1, 99):
        if i not in abs_list:
            return i

def set_numbering(doc, config:HiddenCleanerConfig):
    # 检查是否需要删除编号
    if config.extend.设置标题编号.get():
        标题编号 = [config.extend.一级编号.get(), config.extend.二级编号.get(), config.extend.三级编号.get(), config.extend.四级编号.get(), config.extend.五级编号.get(), config.extend.六级编号.get(), config.extend.七级编号.get(), config.extend.八级编号.get(), config.extend.九级编号.get()]
        标题Fmt = [fmt_dict[config.extend.一级编号Fmt.get()], fmt_dict[config.extend.二级编号Fmt.get()], fmt_dict[config.extend.三级编号Fmt.get()], fmt_dict[config.extend.四级编号Fmt.get()], fmt_dict[config.extend.五级编号Fmt.get()], fmt_dict[config.extend.六级编号Fmt.get()], fmt_dict[config.extend.七级编号Fmt.get()], fmt_dict[config.extend.八级编号Fmt.get()], fmt_dict[config.extend.九级编号Fmt.get()]]
        标题Lgl = [config.extend.一级编号Lgl.get(), config.extend.二级编号Lgl.get(), config.extend.三级编号Lgl.get(), config.extend.四级编号Lgl.get(), config.extend.五级编号Lgl.get(), config.extend.六级编号Lgl.get(), config.extend.七级编号Lgl.get(), config.extend.八级编号Lgl.get(), config.extend.九级编号Lgl.get()]
        try:
            # 使用numbering_part
            numXML = doc.part.numbering_part.numbering_definitions._numbering
            abstractId = -1
            abstractIds = []
            for num in numXML.num_lst:
                if num.numId == config.extend.标题编号ID.get():
                    abstractId = num.abstractNumId.val
                abstractIds.append(num.abstractNumId.val)
            if abstractId>=0:
                # 如果有的话，直接修改
                for absNum in numXML.iter(qn("w:abstractNum")):
                    if int(absNum.get(qn("w:abstractNumId"))) == abstractId:
                        for lvl in absNum.iter(qn("w:lvl")):
                            level = int(lvl.get(qn("w:ilvl")))
                            # 内容
                            lvlText = lvl.find(qn("w:lvlText"))
                            lvlText.set(qn("w:val"), 标题编号[level])
                            # 起始编号
                            start = lvl.find(qn("w:start"))
                            if start is not None:
                                start.set(qn("w:val"), '1')
                            # numFmt
                            numFmt = lvl.find(qn("w:numFmt"))
                            if numFmt is not None:
                                numFmt.set(qn("w:val"), 标题Fmt[level])
                            # suff
                            suff = lvl.find(qn("w:suff"))
                            if suff is not None:
                                suff.set(qn("w:val"), "nothing")
                            else:
                                add_xml(lvl, "w:suff", "nothing")
                            # 正规编号
                            if 标题Lgl[level]:
                                isLgl = lvl.find(qn("w:isLgl"))
                                if not isLgl:
                                    add_xml(lvl, "w:isLgl")
                            else:
                                isLgl = lvl.find(qn("w:isLgl"))
                                if isLgl:
                                    isLgl.getparent().remove(isLgl)
                            # 字体：删除字体，编号字体与正文字体保持一致
                            rPr = lvl.find(qn("w:rPr"))
                            rPr.getparent().remove(rPr)
                            # 缩进
                            pPr = lvl.find(qn("w:pPr"))
                            tabs = pPr.find(qn("w:tabs"))
                            if tabs is not None:
                                tab = tabs.find(qn("w:tab"))
                                if tab is not None:
                                    tab.set(qn("w:val"), "left")
                                    tab.set(qn("w:pos"), '0')
                                else:
                                    add_xml(tabs, "w:tab", {"w:val": "left", "w:pos":"0"})
                            else:
                                tabs = add_xml(pPr, "w:tabs")
                                add_xml(tabs, "w:tab", {"w:val": "left", "w:pos": "0"})
                            ind = pPr.find(qn("w:ind"))
                            if ind is not None:
                                ind.set(qn("w:left"), '0')
                                ind.set(qn("w:hanging"), '0')
                                ind.set(qn("w:firstLine"), '0')
                                ind.set(qn("w:leftChars"), '0')
                                ind.set(qn("w:firstLineChars"), '0')
                            else:
                                add_xml(pPr, "w:ind", {"w:left": "0", "w:hanging": "0", "w:firstLine": "0", "w:leftChars": "0", "w:firstLineChars": "0"})
            else:
                # 如果没有的话，创建新的
                # TODO:考虑的似乎过于简单
                # 1、创建numid和abstracNumId的映射
                abstractId = get_abstract_num(abstractIds)
                abstract_xml = OxmlElement('w:abstractNumId')
                abstract_xml.set(qn('w:val'), abstractId)
                wnum_xml = OxmlElement('w:num')
                wnum_xml.set(qn('w:numId'), config.extend.标题编号ID.get())
                wnum_xml.append(abstract_xml)
                numXML.append(wnum_xml)
                # 2、创建abstractNum
                abstractNum = OxmlElement('w:abstractNum')
                abstractNum.set(qn('w:abstractNumId'), abstractId)
                for i in range(9):
                    lvl = OxmlElement('w:lvl')
                    lvl.set(qn('w:ilvl'), str(i))
                    lvlText = OxmlElement('w:lvlText')
                    lvlText.set(qn('w:val'), 标题编号[i])
                    lvl.append(lvlText)
                    abstractNum.append(lvl)
                numXML.append(abstractNum)
            return True
        except:
            global fmt_index
            fmt_index = [0,0,0,0,0,0,0,0,0]
            # 对于未定义编号的word，不存在numbering，使用文本替换方式
            return False
    else:
        return False


def chinese_counting(num):
    """
    数字转为汉字
    """
    if num < 1:
        return ""

    # 定义数字和汉字的映射
    digits = "零一二三四五六七八九"
    units = ["", "十", "百", "千", "万", "亿"]

    # 递归函数，将数字转换为汉字
    def convert(num, unit_index):
        if num == 0:
            return ""

        # 获取当前位的数字和单位
        digit = num % 10
        unit = units[unit_index]

        # 递归处理下一位
        rest = num // 10
        rest_chinese = convert(rest, unit_index + 1)

        # 如果当前位是0，但下一位不是0，则需要加上"零"
        if digit == 0 and rest_chinese:
            return "零" + rest_chinese
        else:
            return digits[digit] + unit + rest_chinese

    # 调用递归函数
    return convert(num, 0)


def get_str_title(fmt, number):
    if fmt == "decimal":
        return str(number)
    elif fmt == "chineseCounting":
        return chinese_counting(number)
    else:
        return str(number)

def del_paragraph_numbering(paragraph, config:HiddenCleanerConfig):
    level_element = paragraph.style.element.pPr.find(qn("w:outlineLvl"))
    if level_element is not None:
        # 删除原来的编号
        pattern = config.extend.原有标题编号样式.get().split(";")
        sentence = paragraph.text
        for p in pattern:
            match = re.match(p, sentence)
            if match is not None:
                paragraph.text = sentence[match.end():]

def set_paragraph_numbering(paragraph, config:HiddenCleanerConfig):
    level_element = paragraph.style.element.pPr.find(qn("w:outlineLvl"))
    if level_element is not None:
        # 设置新的编号
        outline_level = int(level_element.get(qn("w:val")))
        global fmt_index
        for index in range(outline_level):
            if fmt_index[index] == 0:
                fmt_index[index] = 1
        fmt_index[outline_level] += 1
        for index in range(outline_level+1, 9):
            fmt_index[index] = 0
        标题编号 = [config.extend.一级编号.get(), config.extend.二级编号.get(), config.extend.三级编号.get(), config.extend.四级编号.get(), config.extend.五级编号.get(), config.extend.六级编号.get(), config.extend.七级编号.get(), config.extend.八级编号.get(), config.extend.九级编号.get()]
        标题Fmt = [fmt_dict[config.extend.一级编号Fmt.get()], fmt_dict[config.extend.二级编号Fmt.get()], fmt_dict[config.extend.三级编号Fmt.get()], fmt_dict[config.extend.四级编号Fmt.get()], fmt_dict[config.extend.五级编号Fmt.get()], fmt_dict[config.extend.六级编号Fmt.get()], fmt_dict[config.extend.七级编号Fmt.get()], fmt_dict[config.extend.八级编号Fmt.get()], fmt_dict[config.extend.九级编号Fmt.get()]]
        标题Lgl = [config.extend.一级编号Lgl.get(), config.extend.二级编号Lgl.get(), config.extend.三级编号Lgl.get(), config.extend.四级编号Lgl.get(), config.extend.五级编号Lgl.get(), config.extend.六级编号Lgl.get(), config.extend.七级编号Lgl.get(), config.extend.八级编号Lgl.get(), config.extend.九级编号Lgl.get()]
        if 标题Lgl[outline_level]:
            result = 标题编号[outline_level]
            for i in range(outline_level+1):
                j = i + 1
                result = result.replace(f"%{j}", str(fmt_index[i]))
        else:
            result = 标题编号[outline_level]
            for i in range(outline_level+1):
                j = i + 1
                result = result.replace(f"%{j}", get_str_title(标题Fmt[i], fmt_index[i]))
        paragraph.text = result + paragraph.text

    else:
        return


def set_section(section, config:HiddenCleanerConfig):
    # 设置页面布局
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_height = Cm(29.7)  # A4纸的高度
    section.page_width = Cm(21.0)  # A4纸的宽度
    # 设置页边距
    section.top_margin = Cm(config.page.页边距.上.get())  # 上边距
    section.bottom_margin = Cm(config.page.页边距.下.get())  # 下边距
    section.left_margin = Cm(config.page.页边距.左.get())  # 左边距
    section.right_margin = Cm(config.page.页边距.右.get())  # 右边距
    # 页眉页脚
    if config.core.删除页眉页脚:
        header = section.header
        header.is_linked_to_previous = True
        footer = section.footer
        footer.is_linked_to_previous = True

def set_core(doc, config:HiddenCleanerConfig):
    core_properties = doc.core_properties
    if config.core.删除文档属性.get():
        core_properties.author = ""
        core_properties.category = ""
        core_properties.comments = ""
        core_properties.title = ""
        core_properties.company = ""

def get_bodys(doc):
    body = doc._body._body
    # ps = body.xpath('//w:p')
    ps = body.xpath('.//w:p[not(ancestor::w:tbl)]') # 排除掉表格中的内容
    for p in ps:
        yield Paragraph(p, doc._body)

def set_docx_one(config:HiddenCleanerConfig, file, process=None, count=None, tip=None):
    try:
        doc = Document(file)
    except:
        return "文件读取失败！"

    nowstep = 0
    mainstep = count*0.1
    # 作者等
    if process:
        tip.set("正在处理文件:%s  处理文档信息"%file.split("/")[-1])
    set_core(doc, config)
    if process:
        process.step(mainstep)
        process.update()
        nowstep += mainstep
    # 页面布局
    sections = doc.sections
    if process:
        tip.set("正在处理文件:%s  处理页面布局"%file.split("/")[-1])
    mainstep = count*0.1
    s_step = mainstep/len(sections) if len(sections)>0 else mainstep
    for section in sections:
        set_section(section, config)
        if process:
            process.step(s_step)
            process.update()
            nowstep += s_step
    # 样式
    for style in doc.styles:
        set_style(style, config)

    # 标题编号
    have_numbering = set_numbering(doc, config)
    # 封面或目录
    if config.extend.封面目录处理.get():
        paragraphs = list(get_bodys(doc))
    else:
        paragraphs = doc.paragraphs
    if process:
        tip.set("正在处理文件:%s  处理正文"%file.split("/")[-1])
    mainstep = count * 0.6
    if len(paragraphs) > 0:
        m_step = mainstep/len(paragraphs)
    else:
        process.step(mainstep)
        process.update()
        nowstep += mainstep
    # 正文
    for paragraph in paragraphs:
        set_paragraph(paragraph, config, table=False)
        if config.extend.设置标题编号.get() and config.extend.删除原有标题编号.get():
            del_paragraph_numbering(paragraph, config)
        for run in paragraph.runs:
            set_font(run, config.main.font, config.base)
        if config.extend.设置标题编号.get() and not have_numbering:
            set_paragraph_numbering(paragraph, config)

        # 如果删除图片有可能导致有空行
        if config.base.删除图片.get() and config.base.删除空行.get() and paragraph.text.strip() == "":
            try:
                paragraph._element.getparent().remove(paragraph._element)
            except:
                pass
        if process:
            process.step(m_step)
            process.update()
            nowstep += m_step
    # 表格
    if process:
        tip.set("正在处理文件:%s  处理表格"%file.split("/")[-1])
    mainstep = count * 0.2
    if len(doc.tables) > 0:
        t_step = mainstep/len(doc.tables)
    else:
        process.step(mainstep)
        process.update()
        nowstep += mainstep
    for table in doc.tables:
        set_table(table, config)
        for row in table.rows:
            set_row(row, config)
            for cell in row.cells:
                set_cell(cell, config)
                for paragraph in cell.paragraphs:
                    # set_table_paragraph(paragraph, config)
                    set_paragraph(paragraph, config, table=True)
                    for run in paragraph.runs:
                        # set_table_font(run, config)
                        set_font(run, config.table.font, config.base)
        if process:
            process.step(t_step)
            process.update()
            nowstep += t_step
    # 保存
    if process:
        tip.set("正在处理文件:%s  文件保存" % file.split("/")[-1])
    try:
        doc.save(file.replace('.docx', '_new.docx'))
    except:
        return "文件保存失败！"
    # 在资源管理器中定位到文件
    dir = file.replace('.docx', '_new.docx').replace("/", '\\')
    os.system(r"explorer.exe /select, %s" % dir)
    return True

def set_docx(config:HiddenCleanerConfig, process, tip):
    files = config.file.get().split(";")
    if len(files)==0 or (len(files)==1 and files[0]==""):
        return "没有文件"
    all = len(files)
    if all==1:
        procount = [100]
    else:
        procount = [100//all for _ in range(all-1)]
        procount.append(100-sum(procount))

    index = 0
    for file in files:
        tip.set("正在处理文件:%s"%file.split("/")[-1])
        if check_docx(file):
            result = set_docx_one(config, file, process, procount[index], tip)
            if type(result)==str:
                tip.set("保存失败，请在wps/office中关闭文件并重试！")
                return result
        else:
            return "文件不存在！"
    tip.set("完成！")
    return True




if __name__ == '__main__':
    pass