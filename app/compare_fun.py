import time
import re
import os
import docx as dx
from app.compare_config import CompareConfig
from docx.shared import Pt, RGBColor, Cm, Inches

# 获取文档中的所有段落
def get_paragraphs_from_doc(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


# 获取段落中的所有短句
def get_sentence_from_paragraph(paragraph, config:CompareConfig):
    splitword = config.splitword.get()
    limitnum = int(config.splitnum.get())

    text = paragraph.text
    text = text.replace(" ", "").replace("\t", "")
    if len(splitword)>0:
        for t in re.split(splitword, text):
            if len(t) >= limitnum:
                print(t)
                yield t
    else:
        if len(text) >= limitnum:
            yield text

# 获取文档中的所有短句
def get_sentence(doc, config:CompareConfig):
    for p in get_paragraphs_from_doc(doc):
        for t in get_sentence_from_paragraph(p, config):
            yield t


# 给段落染色
def set_color(paragraph, color):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(color[0],color[1], color[2])



def compare(config:CompareConfig, process, tip):
    doc_orgs = config.zbfiles.get().split(";")
    files = config.tbfiles.get().split(";")
    # 获取招标文件所有字句
    if process:
        tip.set("读取招标文件……")
    text_org = set()
    doc_org_all = list()
    for docfile in doc_orgs:
        doc = dx.Document(docfile)
        doc_org_all.append(doc)
        for t in get_sentence(doc, config):
            text_org.add(t)
    if process:
        process.step(10)
        process.update()
        tip.set("读取投标文件……")
    # 获取招标文件中连续文本
    # senten = ""
    # for doc in doc_org_all:
    #     for text in get_sentence(doc, config):
    #         senten = senten + text
    # 获取文档内容
    text_list = list()
    doc_all = list()
    for i in range(len(files)):
        text_set = set()
        doc = dx.Document(files[i])
        doc_all.append(doc)
        for t in get_sentence(doc, config):
            text_set.add(t)
        text_list.append(text_set)
    if process:
        process.step(15)
        process.update()
        tip.set("开始比对……")
    # 寻找重复字句
    text_error_list = []
    for i in range(len(doc_all)):
        if process:
            tip.set("处理第{}个文件……".format(i+1))
        doc = doc_all[i]
        text_error = set()
        for paragraph in get_paragraphs_from_doc(doc):
            texts = list(get_sentence_from_paragraph(paragraph, config))
            # 删除所有在招标文件中出现的短句
            text_inorg = []
            for text in texts:
                if text in text_org:
                    text_inorg.append(text)
            for text in text_inorg:
                texts.remove(text)
            # 依次与其他文档比较重复
            repeat = []
            for j in range(len(text_list)):
                if j!=i:
                    clone_count = 0
                    for text in texts:
                        if text in text_list[j]:
                            text_error.add(text)
                            clone_count += 1
                    repeat.append(clone_count)
                else:
                    repeat.append(0)
            # 检查重复是否达到阈值
            maxrepeat = max(repeat)
            if maxrepeat == 0:
                pass
            else:
                repeat_num = float(config.repeatnum.get())
                if repeat_num>1:
                    if maxrepeat >= repeat_num:
                        index = repeat.index(maxrepeat)
                        set_color(paragraph, config.tbcolor[index])
                else:
                    if maxrepeat >= repeat_num * len(texts):
                        index = repeat.index(maxrepeat)
                        set_color(paragraph, config.tbcolor[index])

        text_error_list.append(text_error)

        try:
            doc.save(files[i].replace(".docx", "_输出.docx"))
        except PermissionError:
            if process:
                tip.set("保存第%s个文件……拒绝访问，请确认是否已在office中关闭此文件【%s】。将在15s后重新尝试保存。"%(i+1, files[i].replace(".docx", "_输出.docx")))
            time.sleep(30)
            try:
                doc.save(files[i].replace(".docx", "_输出.docx"))
            except:
                return "保存失败，请在wps/office中关闭文件并重试！"
        with open(files[i].replace(".docx", "_重复子句.txt"), 'w') as f:
            for text in text_error_list[i]:
                try:
                    f.write(text + '\n')
                except:
                    pass
        if process:
            process.step(75/len(files))
            process.update()

    tip.set("【处理完成！】")
    # 在资源管理器中定位到文件
    dir = files[0].replace('.docx', '_输出.docx').replace("/", '\\')
    os.system(r"explorer.exe /select, %s" % dir)
    return True



